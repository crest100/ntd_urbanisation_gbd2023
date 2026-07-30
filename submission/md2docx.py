from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re, os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

def add_heading_styled(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

with open(os.path.join(os.path.dirname(__file__), 'manuscript_BMJGH.md'), 'r', encoding='utf-8') as f:
    lines = f.readlines()

in_refs = False
ref_counter = 0
table_capture = False
table_rows = []
in_figure = False

for line in lines:
    stripped = line.rstrip()

    # Skip image-only lines (handled below)
    if stripped.startswith('!['):
        continue

    # Horizontal rule
    if stripped == '---':
        continue

    # References section
    if stripped.startswith('## References'):
        in_refs = True
        doc.add_page_break()
        add_heading_styled('References', level=1)
        continue

    if in_refs:
        if stripped == '':
            doc.add_paragraph('')
        else:
            m = re.match(r'^(\d+)\.\s+(.*)', stripped)
            if m:
                ref_counter += 1
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Inches(-0.5)
                p.paragraph_format.left_indent = Inches(0.5)
                p.add_run(f'{m.group(1)}. ').bold = False
                p.add_run(m.group(2))
            else:
                doc.add_paragraph(stripped)
        continue

    # Headings
    if stripped.startswith('# '):
        add_heading_styled(stripped[2:], level=1)
    elif stripped.startswith('## '):
        add_heading_styled(stripped[3:], level=2)
    elif stripped.startswith('### '):
        add_heading_styled(stripped[4:], level=3)
    elif stripped.startswith('#### '):
        add_heading_styled(stripped[5:], level=4)

    # Figure captions and images
    elif stripped.startswith('!['):
        m = re.match(r'!\[(.*?)\]\((.+?)\)', stripped)
        if m:
            caption = m.group(1)
            img_path = os.path.join(os.path.dirname(__file__), m.group(2))
            # Center the image
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if os.path.exists(img_path):
                run = p.add_run()
                run.add_picture(img_path, width=Inches(5.5))
            # Add caption below
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_before = Pt(4)
            r2 = p2.add_run(caption)
            r2.italic = True
            r2.font.size = Pt(10)

    # Tables
    elif stripped.startswith('|') and stripped.endswith('|'):
        if not table_capture:
            table_capture = True
            table_rows = [stripped]
        else:
            table_rows.append(stripped)
    else:
        if table_capture:
            table_capture = False
            # Write the table
            parsed = []
            for tr in table_rows:
                cells = [c.strip() for c in tr.split('|')[1:-1]]
                parsed.append(cells)
            if len(parsed) >= 3:
                # Second row is the separator; skip it
                header = parsed[0]
                data = parsed[2:]
                t = doc.add_table(rows=1 + len(data), cols=len(header))
                t.style = 'Light Grid Accent 1'
                t.alignment = WD_TABLE_ALIGNMENT.CENTER
                for j, h in enumerate(header):
                    cell = t.rows[0].cells[j]
                    cell.text = h
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(9)
                for i, row in enumerate(data):
                    for j, val in enumerate(row):
                        cell = t.rows[i+1].cells[j]
                        cell.text = val
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
            doc.add_paragraph('')
            table_rows = []
            continue

        # Regular text: handle bold markers
        if stripped == '':
            doc.add_paragraph('')
        else:
            p = doc.add_paragraph()
            # Handle inline bold **...**
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

# Save
output_path = os.path.join(os.path.dirname(__file__), 'Manuscript_BMJGH.docx')
doc.save(output_path)
print(f'Word document saved to: {output_path}')
